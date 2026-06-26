import datetime
import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO # Untuk membuat file Word sementara di memori.

from models import (
    init_db,
    insert_logbook,
    get_all_logbooks,
    get_logbook_by_id,
    update_logbook,
    delete_logbook,
    get_total_duration
)

init_db()

# Perhitungan Durasi Belajar
def format_duration(duration_seconds):
    hour = duration_seconds // 3600
    minute = (duration_seconds % 3600) // 60
    second = duration_seconds % 60
    
    return f"{hour:02}:{minute:02}:{second:02}"

# Generate Word
def generate_word_report(logs):
    doc = Document()

    doc.add_heading(
        "Capstone Project Logbook",
        level=1
    )

    table = doc.add_table(
        rows=1,
        cols=5
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = "Tanggal"
    header[1].text = "Durasi"
    header[2].text = "Tugas/Proyek"
    header[3].text = "Hasil dan Kendala"
    header[4].text = "Solusi / Action Plan"

    for log in logs:
        row = table.add_row().cells

        row[0].text = str(log[1])
        row[1].text = format_duration(log[2])
        row[2].text = str(log[3])
        row[3].text = str(log[4])
        row[4].text = str(log[5])

    file_stream = BytesIO()

    doc.save(file_stream) # Simpan dokumen Word ke memori ini.

    file_stream.seek(0) # Kembali ke awal file supaya bisa dibaca/download dari awal.

    return file_stream

st.set_page_config(
    page_title="Capstone Logbook Generator",
    layout="wide"
)

st.title("📑 Capstone Logbook Generator 📑")

if "start_time" not in st.session_state:
    st.session_state.start_time = None

# Waktu Mulai
if st.button("▶️ Mulai Sesi Belajar ▶️"):
    st.session_state.start_time = datetime.datetime.now()

    # st.success(
    #     f"Sesi dimulai pada: "
    #     f"{st.session_state.start_time.strftime('%H:%M:%S')}"
    # )

if st.session_state.start_time is not None:
    st.info(
        f"Sesi dimulai sejak: "
        f"{st.session_state.start_time.strftime('%H:%M:%S')}"
    )

# input("Tekan ENTER saat sesi selesai dan mengisi logbook...")

st.divider()
st.subheader("📝 Isi Laporan Logbook 📝")

# Input User
activity = st.text_area("Tugas/proyek yang dikerjakan: ")
result = st.text_area("Hasil dan kendala: ")
action_plan = st.text_area("Solusi atau action plan selanjutnya: ")

# Waktu Selesai
if st.button("⏹️ Akhiri Sesi Belajar ⏹️"):
    if st.session_state.start_time is None:
        st.error("Silakan mulai sesi belajar terlebih dahulu.")
    elif not activity.strip():
        st.error("Tugas/Proyek tidak boleh kosong.")
    elif not result.strip():
        st.error("Hasil dan kendala tidak boleh kosong.")
    elif not action_plan.strip():
        st.error("Solusi atau action plan tidak boleh kosong.")
    else:
        finish_time = datetime.datetime.now()
        st.write(
            f"Sesi selesai pada: "
            f"{finish_time.strftime('%H:%M:%S')}"
        )

        # Tanggal
        session_date = (st.session_state.start_time.strftime("%d-%m-%Y"))

        # Perhitungan Durasi Belajar
        duration = (finish_time - st.session_state.start_time)
        duration_seconds = int(duration.total_seconds())
        formatted_duration = format_duration(duration_seconds)

        st.success(
            f"Total durasi sesi belajar: "
            f"{formatted_duration}"
        )

        # Simpan ke database
        insert_logbook(
            session_date,
            duration_seconds,
            activity,
            result,
            action_plan
        )

        st.success(
            "Logbook berhasil disimpan ke database!"
        )

        # Reset sesi belajar.
        st.session_state.start_time = None

st.divider()

st.subheader("📄 Rekapitulasi Laporan Logbook 📄")

logs = get_all_logbooks()

total_session = len(logs)

total_seconds = get_total_duration()

if total_seconds is None:
    total_seconds = 0

formatted_total_duration = format_duration(total_seconds)

col1, col2 = st.columns(2)

with col1:
    st.metric(
        "Total Sesi Belajar",
        total_session
    )

with col2:
    st.metric(
        "Total Durasi Belajar",
        formatted_total_duration
    )

st.divider()

st.subheader("📚 Preview Laporan Logbook 📚")

if logs:
    preview_data = []

    for log in logs:
        preview_data.append({
            "Tanggal": log[1],
            "Durasi": format_duration(log[2]),
            "Tugas/Proyek": log[3],
            "Hasil dan Kendala": log[4],
            "Solusi atau Action Plan": log[5]
        })
    
    preview_df = pd.DataFrame(preview_data)

    st.dataframe(
        preview_df,
        use_container_width=True,
        hide_index=True
    )

else:
    st.info("Belum ada data logbook.")

st.divider()

# Callback Selectbox
def load_logbook():
    selected_id = st.session_state.selected_id

    selected_log = get_logbook_by_id(selected_id)

    if selected_log:
        st.session_state.edit_activity = selected_log[3]
        st.session_state.edit_result = selected_log[4]
        st.session_state.edit_action_plan = selected_log[5]

st.subheader("✏️ Edit & Hapus Logbook ✏️")

if "selected_id" not in st.session_state:
    st.session_state.selected_id = None

if "edit_activity" not in st.session_state:
    st.session_state.edit_activity = ""

if "edit_result" not in st.session_state:
    st.session_state.edit_result = ""

if "edit_action_plan" not in st.session_state:
    st.session_state.edit_action_plan = ""

if logs:
    options = {
        f"ID {log[0]} | {log[1]}": log[0]
        for log in logs
    }

    selected_label = st.selectbox(
        "Pilih Logbook",
        options=list(options.keys())
    )

    selected_id = options[selected_label]

    if st.session_state.selected_id != selected_id:
        st.session_state.selected_id = selected_id
        load_logbook()

    edit_activity = st.text_area(
        "Tugas/Proyek",
        key="edit_activity"
    )

    edit_result = st.text_area(
        "Hasil dan Kendala",
        key="edit_result"
    )

    edit_action_plan = st.text_area(
        "Solusi atau Action Plan",
        key="edit_action_plan"
    )

    col1, col2 = st.columns(2)

    with col1:
        if st.button("💾 Update Logbook"):
            update_logbook(
                selected_id,
                edit_activity,
                edit_result,
                edit_action_plan
            )

            st.success(
                "Logbook berhasil diperbarui!"
            )

            st.rerun()
    
    with col2:
        if st.button("🗑️ Hapus Logbook"):
            delete_logbook(
                selected_id
            )
            
            st.success(
                "Logbook berhasil dihapus!"
            )

            st.rerun()

else:
    st.info(
        "Pilih data logbook untuk diperbarui atau dihapus."
    )

st.divider()

st.subheader("🖨️ Generate Logbook Report 🖨️")

if logs:
    word_file = generate_word_report(logs)

    st.download_button(
        label="⬇️ Download Logbook (.docx)",
        data=word_file,
        file_name="Capstone_Logbook.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.info(
        "Belum ada data logbook untuk di-generate."
    )